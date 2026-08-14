"""Извлечение текста из вложенных документов.

Внедренец получает от заказчика ТЗ, протоколы и коммерческие предложения
в PDF и DOCX. Печатать их содержимое в чат никто не будет, поэтому файл
должен разбираться так же, как напечатанные заметки.
"""

import io
import os

# Сколько символов отдаём модели. Большое ТЗ целиком в контекст не влезет,
# а обрезать молча нельзя — человеку сообщается, что документ прочитан
# не полностью.
MAX_CHARS = 20000

SUPPORTED = (".pdf", ".docx", ".txt", ".md", ".csv")


class UnsupportedFile(Exception):
    pass


def _from_pdf(data):
    try:
        from pypdf import PdfReader
    except ImportError:
        raise UnsupportedFile("не установлен pypdf: pip install pypdf")

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        raise UnsupportedFile(
            "в PDF нет текстового слоя — похоже, это скан. "
            "Нужен OCR, он в прототип не входит")
    return text


def _from_docx(data):
    try:
        import docx
    except ImportError:
        raise UnsupportedFile("не установлен python-docx: pip install python-docx")

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Требования в ТЗ часто лежат в таблицах, а не в абзацах.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _from_text(data):
    for encoding in ("utf-8", "cp1251"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def extract(filename, data):
    """Текст документа. Бросает UnsupportedFile с понятным объяснением."""
    extension = os.path.splitext(filename or "")[1].lower()

    if extension == ".pdf":
        text = _from_pdf(data)
    elif extension == ".docx":
        text = _from_docx(data)
    elif extension in (".txt", ".md", ".csv"):
        text = _from_text(data)
    elif extension == ".doc":
        raise UnsupportedFile(
            "старый формат .doc не поддерживается — пересохраните как .docx")
    else:
        raise UnsupportedFile(
            f"формат {extension or 'без расширения'} не поддерживается. "
            f"Подойдут: {', '.join(SUPPORTED)}")

    if not text:
        raise UnsupportedFile("документ пустой или из него не удалось извлечь текст")

    truncated = len(text) > MAX_CHARS
    return text[:MAX_CHARS], truncated


def as_message(filename, data):
    """Готовый текст сообщения для агента."""
    text, truncated = extract(filename, data)
    header = f"Содержимое файла «{filename}»"
    if truncated:
        header += f" (прочитаны первые {MAX_CHARS} символов, документ длиннее)"
    return f"{header}:\n\n{text}"
